"""Модуль для безопасного чтения файлов с диска."""

import os
from pathlib import Path
from typing import Optional

from kim_core.config.settings import AppConfig
from kim_core.logging import logger


class FileAccessError(Exception):
    """Базовое исключение для ошибок доступа к файлам."""

    pass


class FileTypeNotSupportedError(FileAccessError):
    """Исключение для неподдерживаемых типов файлов."""

    pass


def is_path_allowed(path: Path, whitelist_dirs: list[Path]) -> bool:
    """
    Проверяет, разрешён ли доступ к файлу по белому списку директорий.

    Args:
        path: Путь к файлу
        whitelist_dirs: Список разрешённых директорий

    Returns:
        True, если путь разрешён, False иначе
    """
    if not whitelist_dirs:
        logger.warning("Белый список директорий пуст - доступ запрещён")
        return False

    try:
        # Нормализуем путь (разрешаем символические ссылки и относительные пути)
        resolved_path = path.resolve()
        
        # Проверяем, что файл находится внутри одной из разрешённых директорий
        for whitelist_dir in whitelist_dirs:
            try:
                resolved_whitelist = Path(whitelist_dir).resolve()
                
                # Проверяем, что путь является подпутем разрешённой директории
                # Для Python 3.9+ используем is_relative_to, для совместимости - альтернативный метод
                try:
                    # Python 3.9+
                    if resolved_path.is_relative_to(resolved_whitelist):
                        logger.debug(
                            f"Путь разрешён: {resolved_path} находится в {resolved_whitelist}"
                        )
                        return True
                except AttributeError:
                    # Для Python < 3.9 используем альтернативный метод
                    try:
                        resolved_path.relative_to(resolved_whitelist)
                        logger.debug(
                            f"Путь разрешён: {resolved_path} находится в {resolved_whitelist}"
                        )
                        return True
                    except ValueError:
                        # Не является подпутем
                        continue
                        
            except (OSError, ValueError) as e:
                logger.warning(f"Ошибка при проверке пути {whitelist_dir}: {e}")
                continue

        logger.warning(f"Путь не разрешён: {resolved_path} не находится в белом списке")
        return False

    except (OSError, ValueError) as e:
        logger.error(f"Ошибка при разрешении пути {path}: {e}")
        return False


def truncate_text(text: str, max_chars: int) -> str:
    """
    Обрезает текст до указанной длины.

    Args:
        text: Исходный текст
        max_chars: Максимальная длина в символах

    Returns:
        Обрезанный текст с пометкой о сокращении
    """
    if len(text) <= max_chars:
        return text

    truncated = text[:max_chars]
    return truncated + "... [truncated]"


def read_file_text(path: Path, config: AppConfig) -> str:
    """
    Безопасно читает текстовое содержимое файла.

    Args:
        path: Путь к файлу
        config: Конфигурация приложения

    Returns:
        Текстовое содержимое файла

    Raises:
        FileAccessError: Если доступ к файлу запрещён или файл не существует
        FileTypeNotSupportedError: Если тип файла не поддерживается
    """
    # Проверка существования файла
    if not path.exists():
        error_msg = f"Файл не существует: {path}"
        logger.error(error_msg)
        raise FileAccessError(error_msg)

    # Проверка, что это не директория
    if path.is_dir():
        error_msg = f"Указанный путь является директорией, а не файлом: {path}"
        logger.error(error_msg)
        raise FileAccessError(error_msg)

    # Проверка размера файла
    try:
        file_size = path.stat().st_size
        file_size_mb = file_size / (1024 * 1024)
        
        if file_size_mb > config.file_max_size_mb:
            error_msg = (
                f"Файл слишком большой: {file_size_mb:.2f} MB "
                f"(максимум: {config.file_max_size_mb} MB)"
            )
            logger.error(error_msg)
            raise FileAccessError(error_msg)
            
        logger.info(
            f"Чтение файла: {path}, размер: {file_size_mb:.2f} MB, "
            f"размер в байтах: {file_size}"
        )
    except OSError as e:
        error_msg = f"Ошибка при проверке размера файла {path}: {e}"
        logger.error(error_msg)
        raise FileAccessError(error_msg) from e

    # Проверка белого списка директорий
    whitelist_paths = []
    if config.file_whitelist_dirs:
        for whitelist_dir in config.file_whitelist_dirs:
            try:
                whitelist_paths.append(Path(whitelist_dir))
            except Exception as e:
                logger.warning(f"Некорректный путь в белом списке: {whitelist_dir}, ошибка: {e}")

    if not is_path_allowed(path, whitelist_paths):
        error_msg = f"Доступ к файлу запрещён: {path} не находится в белом списке директорий"
        logger.error(error_msg)
        raise FileAccessError(error_msg)

    # Определение типа файла и чтение
    file_ext = path.suffix.lower()
    
    logger.info(f"Чтение файла типа: {file_ext}, путь: {path}")

    try:
        if file_ext == ".txt" or file_ext == ".md":
            return _read_text_file(path)
        elif file_ext == ".pdf":
            return _read_pdf_file(path)
        elif file_ext == ".docx":
            return _read_docx_file(path)
        else:
            error_msg = (
                f"Тип файла не поддерживается: {file_ext}. "
                f"Поддерживаются: .txt, .md, .pdf, .docx"
            )
            logger.error(error_msg)
            raise FileTypeNotSupportedError(error_msg)
            
    except FileTypeNotSupportedError:
        raise
    except Exception as e:
        error_msg = f"Ошибка при чтении файла {path}: {e}"
        logger.error(error_msg, exc_info=True)
        raise FileAccessError(error_msg) from e


def _read_text_file(path: Path) -> str:
    """
    Читает текстовый файл (txt, md).

    Args:
        path: Путь к файлу

    Returns:
        Содержимое файла как строка
    """
    # Пробуем UTF-8, затем cp1251 (Windows-1251), затем latin-1 как последний fallback
    encodings = ["utf-8", "cp1251", "latin-1"]
    
    for encoding in encodings:
        try:
            with open(path, "r", encoding=encoding, errors="replace") as f:
                content = f.read()
                logger.debug(f"Файл успешно прочитан с кодировкой {encoding}")
                return content
        except UnicodeDecodeError:
            continue
        except Exception as e:
            logger.warning(f"Ошибка при чтении с кодировкой {encoding}: {e}")
            continue

    # Если все кодировки не подошли, пробуем с игнорированием ошибок
    try:
        with open(path, "r", encoding="utf-8", errors="ignore") as f:
            content = f.read()
            logger.warning(f"Файл прочитан с игнорированием ошибок кодировки")
            return content
    except Exception as e:
        error_msg = f"Не удалось прочитать текстовый файл {path}: {e}"
        logger.error(error_msg)
        raise FileAccessError(error_msg) from e


def _read_pdf_file(path: Path) -> str:
    """
    Читает PDF-файл.

    Args:
        path: Путь к файлу

    Returns:
        Извлечённый текст из PDF
    """
    try:
        # Пробуем импортировать pypdf (новая версия PyPDF2)
        try:
            from pypdf import PdfReader
        except ImportError:
            # Fallback на PyPDF2
            try:
                from PyPDF2 import PdfReader
            except ImportError:
                error_msg = (
                    "Библиотека для чтения PDF не установлена. "
                    "Установите: pip install pypdf или pip install PyPDF2"
                )
                logger.error(error_msg)
                raise FileAccessError(error_msg)

        reader = PdfReader(str(path))
        text_parts = []
        
        for page_num, page in enumerate(reader.pages, 1):
            try:
                text = page.extract_text()
                if text:
                    text_parts.append(text)
                logger.debug(f"Прочитана страница {page_num} из PDF")
            except Exception as e:
                logger.warning(f"Ошибка при извлечении текста со страницы {page_num}: {e}")
                continue

        if not text_parts:
            logger.warning(f"Не удалось извлечь текст из PDF: {path}")
            return "Не удалось извлечь текст из PDF-файла."

        full_text = "\n\n".join(text_parts)
        logger.info(f"Извлечено {len(full_text)} символов из PDF, страниц: {len(reader.pages)}")
        return full_text

    except FileAccessError:
        raise
    except Exception as e:
        error_msg = f"Ошибка при чтении PDF-файла {path}: {e}"
        logger.error(error_msg)
        raise FileAccessError(error_msg) from e


def _read_docx_file(path: Path) -> str:
    """
    Читает DOCX-файл.

    Args:
        path: Путь к файлу

    Returns:
        Извлечённый текст из DOCX
    """
    try:
        from docx import Document
    except ImportError:
        error_msg = (
            "Библиотека python-docx не установлена. "
            "Установите: pip install python-docx"
        )
        logger.error(error_msg)
        raise FileAccessError(error_msg)

    try:
        doc = Document(str(path))
        text_parts = []
        
        # Извлекаем текст из всех параграфов
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)
        
        # Извлекаем текст из таблиц (опционально)
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text_parts.append(" | ".join(row_text))

        if not text_parts:
            logger.warning(f"Не удалось извлечь текст из DOCX: {path}")
            return "Не удалось извлечь текст из DOCX-файла."

        full_text = "\n\n".join(text_parts)
        logger.info(f"Извлечено {len(full_text)} символов из DOCX")
        return full_text

    except Exception as e:
        error_msg = f"Ошибка при чтении DOCX-файла {path}: {e}"
        logger.error(error_msg)
        raise FileAccessError(error_msg) from e

